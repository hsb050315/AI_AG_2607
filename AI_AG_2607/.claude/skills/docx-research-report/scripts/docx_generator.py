"""python-docx 기반 범용 Word 문서 생성 템플릿."""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TITLE_SIZE = Pt(40)
BODY_SIZE = Pt(20)
DEFAULT_MAIN_COLOR = "1F4E79"  # 남색 계열


def _hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return RGBColor.from_string(hex_color.upper())


def new_document(main_color=None):
    """새 Document 객체 생성. main_color(hex, 예: '2E86AB')로 문서 전체 강조색을 지정한다."""
    doc = Document()
    doc.main_color = _hex_to_rgb(main_color or DEFAULT_MAIN_COLOR)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return doc


def add_title(doc, text, subtitle=None):
    """문서 제목(및 부제)을 추가한다."""
    heading = doc.add_heading("", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(text)
    run.font.size = TITLE_SIZE
    run.font.color.rgb = doc.main_color
    run.font.name = "맑은 고딕"

    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return doc


def add_heading(doc, text, level=1):
    """섹션 제목 추가 (level 1~4). 문서의 main_color를 적용한다."""
    heading = doc.add_heading("", level=level)
    run = heading.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.color.rgb = doc.main_color
    return heading


def add_paragraph(doc, text, bold=False, italic=False, align=None):
    """일반 본문 단락 추가."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p


def add_bullet_list(doc, items):
    """글머리 기호 목록 추가."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc, items):
    """번호 매기기 목록 추가."""
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, style="Light Grid Accent 1"):
    """헤더와 행 데이터로 표 추가. 헤더 행에 main_color 배경을 적용한다."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hex_color = str(doc.main_color)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        _set_cell_shading(hdr_cells[i], hex_color)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    return table


def add_image(doc, path, width_cm=None):
    """이미지 삽입 (width_cm 미지정 시 원본 크기)."""
    if width_cm:
        doc.add_picture(path, width=Cm(width_cm))
    else:
        doc.add_picture(path)


def add_page_break(doc):
    doc.add_page_break()


def add_divider(doc, color=None):
    """섹션 구분용 가로선 추가."""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color or str(doc.main_color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_hyperlink(paragraph, text, url):
    """단락에 하이퍼링크 텍스트를 삽입한다."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(underline)
    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_margins(doc, top_cm=2.5, bottom_cm=2.5, left_cm=2.5, right_cm=2.5):
    """문서 상하좌우 여백 설정."""
    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)


def add_page_number_footer(doc):
    """꼬리말에 자동 페이지 번호를 삽입한다."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    return p


def save(doc, path):
    parent = os.path.dirname(path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(path)
    print(f"저장 완료: {path}")


def build_sample_document(output_path="sample_output.docx", main_color=None):
    """템플릿 기능을 확인할 수 있는 예시 문서 생성. main_color(hex)로 강조색을 바꿀 수 있다."""
    doc = new_document(main_color=main_color)
    set_margins(doc)
    add_page_number_footer(doc)

    add_title(doc, "문서 제목", subtitle="부제 / 작성일 등")

    add_heading(doc, "1. 개요", level=1)
    add_paragraph(doc, "이 문단은 python-docx로 생성된 예시 본문입니다.")
    p = add_paragraph(doc, "관련 링크: ")
    add_hyperlink(p, "python-docx 문서", "https://python-docx.readthedocs.io/")

    add_divider(doc)

    add_heading(doc, "2. 목록 예시", level=1)
    add_bullet_list(doc, ["항목 1", "항목 2", "항목 3"])
    add_numbered_list(doc, ["첫 번째 단계", "두 번째 단계"])

    add_heading(doc, "3. 표 예시", level=1)
    add_table(
        doc,
        headers=["항목", "값", "비고"],
        rows=[
            ["A", "10", "-"],
            ["B", "20", "-"],
        ],
    )

    add_page_break(doc)
    add_heading(doc, "4. 결론", level=1)
    add_paragraph(doc, "필요한 함수를 조합해 원하는 문서를 구성하세요.")

    save(doc, output_path)


if __name__ == "__main__":
    build_sample_document(main_color="2E86AB")
